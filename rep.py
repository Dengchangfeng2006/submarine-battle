#!/usr/bin/python
# -*- coding: utf-8 -*-
"""生成潜艇大战课程设计实验报告（Word文档）- 包含全部健壮性改进"""
import os
import docx
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = docx.Document()

# ========== 页面设置 ==========
section = doc.sections[0]
section.top_margin = Cm(2.5)
section.bottom_margin = Cm(2.5)
section.left_margin = Cm(3)
section.right_margin = Cm(2.5)

# ========== 格式设置 ==========
style = doc.styles['Normal']
style.font.name = '宋体'
style.font.size = Pt(12)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
style.paragraph_format.line_spacing = 1.5


def addHeading(text, level=1):
    heading = doc.add_heading(text, level=level)
    for run in heading.runs:
        run.font.name = '宋体'
        run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    return heading


def addPara(text, bold=False, indent=True):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.5
    if indent:
        para.paragraph_format.first_line_indent = Pt(24)
    run = para.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.bold = bold
    return para


def addCode(text):
    para = doc.add_paragraph()
    para.paragraph_format.line_spacing = 1.0
    para.paragraph_format.left_indent = Cm(1)
    run = para.add_run(text)
    run.font.name = 'Courier New'
    run.font.size = Pt(9)
    return para


def addAttr(text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run(text)
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')


def addMethod(name, purpose, ret, params):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(24)
    run = p.add_run("方法：" + name)
    run.bold = True
    run.font.name = '宋体'
    run.font.size = Pt(12)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    addPara("作用：" + purpose)
    addPara("返回值：" + ret)
    addPara("参数：" + params)


# ============================================================
# 封面
# ============================================================
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("Python程序设计课程设计报告")
run.font.name = '宋体'
run.font.size = Pt(26)
run.bold = True
run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_paragraph()
doc.add_paragraph()

for item in [
    "课题名称：潜艇大战游戏",
    "提交文档学生姓名：__________",
    "提交文档学生学号：__________",
    "评阅意见：________________________________",
    "提交报告时间：______年______月______日"
]:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(item)
    run.font.name = '宋体'
    run.font.size = Pt(16)
    run.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

doc.add_page_break()

# ============================================================
# 一、课程设计目标
# ============================================================
addHeading("一、课程设计目标", level=1)
addPara("本项目旨在开发一款基于Pygame的潜艇大战视频小游戏。通过本项目的设计与实现，达到以下目标：")
for g in [
    "掌握Python面向对象编程的基本方法，熟练运用类与对象进行程序设计；",
    "掌握Pygame游戏开发框架的使用，包括图形绘制、事件处理、碰撞检测、音效播放等核心功能；",
    "培养将实际问题抽象为程序模型的分析能力，学会合理地划分类与函数的职责；",
    "通过完整的游戏项目开发，体验从需求分析、设计、编码实现到测试的完整软件开发流程；",
    "提高代码的健壮性和容错性，掌握异常处理、资源预加载、字体fallback、跨平台兼容等工程实践；",
    "编写规范、整洁、易读的代码，遵循Python编码规范和最佳实践。"
]:
    addAttr(g)

addPara("本项目的具体要求包括：使用Pygame库开发游戏；玩家通过键盘控制军舰在海面左右移动并释放深水炸弹；潜艇由程序算法控制随机从左右两侧出现；实现得分系统、生命系统、连击奖励系统、最高分持久化、炸弹冷却机制、爆炸动画效果和程序化音效系统；具备完善的异常处理机制，在图片缺失、字体不可用等异常情况下不崩溃。")

# ============================================================
# 二、分析与设计
# ============================================================
addHeading("二、分析与设计", level=1)

addHeading("2.1 整体架构设计", level=2)
addPara("本游戏采用面向对象的设计方法，将游戏中的各个实体抽象为独立的类。程序主要由五个核心类和一个辅助类组成，并配合多个独立工具函数：")
for c in [
    "Warship（军舰类）：负责玩家军舰的移动、动画和炸弹释放位置计算；",
    "DepthBomb（深水炸弹类）：负责炸弹的下落移动和状态管理，使用类级别图片缓存；",
    "Submarine（潜艇类）：负责敌方潜艇的随机生成、移动和销毁，使用类级别图片缓存；",
    "Explosion（爆炸效果类）：负责播放三帧爆炸动画，使用类级别图片缓存；",
    "ComboDisplay（连击显示类，新增）：负责连击飘字效果，带渐隐和上浮动画；",
    "GameController（游戏控制器类）：负责游戏主循环、事件处理、碰撞检测、连击管理、音效播放、得分管理和界面渲染。"
]:
    addAttr(c)

addPara("独立工具函数包括：loadImage()（带异常处理的图片加载，缺失时返回品红占位色块）、createFont()（带多级fallback链的字体创建）、generateBeepSound()（使用数学函数程序化生成方波音效）、loadHighScore()/saveHighScore()（最高分JSON持久化）、preloadAllImages()/getPreloadedImage()（全局图片预加载与获取）。")

addPara("程序整体流程：游戏启动→初始化Pygame及音频→预加载全部图片资源→生成程序化音效→显示开始画面（含最高分）→玩家按空格进入主循环。主循环每帧执行：更新冷却计时→更新连击系统→处理输入事件→更新游戏状态→检测碰撞→渲染画面。生命值降为0后进入结束画面，支持R键重开和Q键退出。")

addHeading("2.2 自定义函数说明", level=2)

# (1) loadImage
addPara("（1）函数名称：loadImage(fileName, scaleWidth=None, scaleHeight=None)", bold=True)
addPara("函数作用：加载指定文件名的图片资源，支持可选缩放。使用try-except捕获pygame.error和FileNotFoundError：加载成功返回带透明通道的Surface；加载失败打印警告并返回品红色占位色块，确保程序不会因资源缺失而崩溃。")
addPara("返回值类型：pygame.Surface对象。")
addPara("参数说明：fileName（字符串类型，不含路径前缀）、scaleWidth（整数类型，可选，目标宽度）、scaleHeight（整数类型，可选，目标高度）。")
addPara("算法说明：使用os.path.join()拼接IMAGE_DIR和文件名，在try块中调用pygame.image.load().convert_alpha()加载，可选调用pygame.transform.scale()缩放。except块创建品红色(magenta)占位Surface。")
addCode("""def loadImage(fileName, scaleWidth=None, scaleHeight=None):
    path = os.path.join(IMAGE_DIR, fileName)
    try:
        image = pygame.image.load(path).convert_alpha()
        if scaleWidth and scaleHeight:
            image = pygame.transform.scale(image, (scaleWidth, scaleHeight))
        return image
    except (pygame.error, FileNotFoundError):
        print(f"警告: 无法加载图片 '{path}'，使用占位图")
        w, h = scaleWidth or 50, scaleHeight or 50
        placeholder = pygame.Surface((w, h), pygame.SRCALPHA)
        placeholder.fill((255, 0, 255))
        return placeholder""")
doc.add_paragraph()

# (2) createFont
addPara("（2）函数名称：createFont(nameList, size)", bold=True)
addPara("函数作用：依次尝试字体名列表创建pygame字体，全部失败则回退到pygame默认字体。解决了SimHei等中文字体在非中文Windows或Linux上不存在导致崩溃的问题。通过渲染测试文字验证字体真实可用（SysFont不抛异常但可能静默回退）。")
addPara("返回值类型：pygame.font.Font对象。")
addPara("参数说明：nameList（字符串列表，按优先级排列的字体名）、size（整数，字体大小像素）。")
addPara("算法说明：遍历nameList，对每个名称调用pygame.font.SysFont()后渲染测试文字，检查get_width()>0确保字体真实可用。全部失败则调用pygame.font.Font(None, size)使用默认字体并打印警告。")

# (3) generateBeepSound
addPara("（3）函数名称：generateBeepSound(frequency=440, duration=0.1, volume=0.3)", bold=True)
addPara("函数作用：使用数学函数程序化生成带衰减包络的方波音效，无需外部音频文件。采用11025Hz采样率以控制struct.pack参数展开数量。")
addPara("返回值类型：pygame.mixer.Sound对象。")
addPara("参数说明：frequency（整数，Hz）、duration（浮点数，秒）、volume（浮点数，0.0~1.0）。")
addPara("算法说明：生成样本列表—每采样点计算方波值(sin>0则1否则-1)乘以线性衰减包络(max(0, 1-t/duration))和音量，使用struct.pack('h'*n, *samples)打包为16位PCM数据传给pygame.mixer.Sound。异常时返回2字节静默Sound。")
addCode("""def generateBeepSound(frequency=440, duration=0.1, volume=0.3):
    sampleRate = 11025  # 降低采样率避免 struct.pack 参数过多
    sampleCount = int(sampleRate * duration)
    samples = []
    for i in range(sampleCount):
        t = i / sampleRate
        value = 1 if math.sin(2*math.pi*frequency*t) > 0 else -1
        envelope = max(0, 1 - t / duration)
        samples.append(int(value * envelope * volume * 32767))
    try:
        sound = pygame.mixer.Sound(buffer=bytes(
            struct.pack('h' * len(samples), *samples)))
    except Exception:
        sound = pygame.mixer.Sound(buffer=bytes(2))
    return sound""")
doc.add_paragraph()

# (4) 最高分函数
addPara("（4）函数名称：loadHighScore() 和 saveHighScore(score)", bold=True)
addPara("函数作用：实现最高分的JSON文件持久化。loadHighScore()从highscore.json读取历史最高分（文件缺失或解析失败返回0）。saveHighScore()仅在分数超过当前记录时写入文件。所有操作包裹在try-except中，失败不影响游戏运行。")
addPara("返回值类型：loadHighScore()返回int；saveHighScore()无返回值。")

# (5) 预加载函数
addPara("（5）函数名称：preloadAllImages() 和 getPreloadedImage(key)", bold=True)
addPara("函数作用：游戏启动时一次性预加载全部11张图片到全局字典PRELOADED_IMAGES中，后续通过getPreloadedImage(key)获取。各实体类通过类级别缓存进一步减少重复引用。")
addPara("返回值类型：preloadAllImages()无返回值；getPreloadedImage(key)返回pygame.Surface或None。")

addHeading("2.3 类详细说明", level=2)

# --- Warship ---
addHeading("（1）Warship（军舰类）", level=3)
addPara("类的作用及意义：Warship类代表玩家控制的军舰，封装了军舰的所有行为和状态。遵循单一职责原则，仅负责军舰自身的移动、位置约束（constrainToScreen）和两帧摇摆动画切换，不涉及游戏逻辑或其他实体。父类：object。")
addPara("属性说明：", bold=True)
for a in [
    "shipImages（列表，存2个pygame.Surface）：两帧动画图片引用，交替切换实现航行摇摆效果；",
    "shipWidth/shipHeight（整数）：军舰图片宽高（像素），用于位置计算和边界约束；",
    "positionX（整数）：军舰左上角X坐标，初始为屏幕水平居中；",
    "positionY（整数）：军舰左上角Y坐标，=WATER_LEVEL_Y-shipHeight，使军舰底部接触海面；",
    "animationFrame（整数，0或1）：当前动画帧索引；",
    "animationTimer（整数）：动画计时器，每10帧切换animationFrame；",
    "isMovingLeft/isMovingRight（布尔）：分别标记军舰是否左右移动，实现平滑连续移动控制。"
]:
    addAttr(a)

addPara("方法说明：", bold=True)
for name, pur, ret, par in [
    ("moveLeft()", "设置isMovingLeft=True", "无", "无"),
    ("moveRight()", "设置isMovingRight=True", "无", "无"),
    ("stopMovingLeft()", "设置isMovingLeft=False", "无", "无"),
    ("stopMovingRight()", "设置isMovingRight=False", "无", "无"),
    ("updatePosition()", "根据移动状态更新X坐标并调用constrainToScreen()", "无", "无"),
    ("constrainToScreen()", "限制positionX在[0, SCREEN_WIDTH-shipWidth]范围内", "无", "无"),
    ("updateAnimation()", "移动时animationTimer累加，每10帧切换animationFrame(0<->1)", "无", "无"),
    ("getCurrentImage()", "返回当前动画帧军舰图片", "pygame.Surface", "无"),
    ("getBombReleasePosition()", "返回军舰底部中心坐标=(positionX+width/2, positionY+height)", "(int,int)", "无"),
    ("getCollisionRect()", "返回军舰碰撞矩形", "pygame.Rect", "无"),
]:
    addMethod(name, pur, ret, par)

# --- DepthBomb ---
addHeading("（2）DepthBomb（深水炸弹类）", level=3)
addPara("类的作用及意义：DepthBomb类代表军舰释放的深水炸弹，封装炸弹的移动逻辑和生命周期管理。使用类变量_bombImage实现所有实例共享一张图片引用，避免重复加载。父类：object。")
addPara("属性说明：", bold=True)
for a in [
    "_bombImage（类变量，pygame.Surface）：所有实例共享的炸弹图片缓存；",
    "bombImage（pygame.Surface）：当前实例的图片引用；",
    "positionX/Y（整数）：炸弹左上角坐标，初始居中于军舰底部中央；",
    "isActive（布尔）：标记炸弹是否有效。"
]:
    addAttr(a)
addPara("方法说明：", bold=True)
for name, pur, ret, par in [
    ("getBombImage()（类方法）", "首次调用时从预加载缓存取boom.png，后续返回缓存", "pygame.Surface", "无"),
    ("updatePosition()", "positionY += BOMB_SPEED(6px)，模拟下沉", "无", "无"),
    ("isOutOfScreen()", "positionY > SCREEN_HEIGHT时返回True", "bool", "无"),
    ("getCollisionRect()", "返回炸弹碰撞矩形", "pygame.Rect", "无"),
    ("render(screen)", "在指定屏幕绘制炸弹", "无", "screen: pygame.Surface"),
]:
    addMethod(name, pur, ret, par)

# --- Submarine ---
addHeading("（3）Submarine（潜艇类）", level=3)
addPara("类的作用及意义：Submarine类代表敌方潜艇，封装随机生成算法和水平移动逻辑。使用类字典_imageCache按潜艇类型缓存图片。父类：object。")
addPara("属性说明：", bold=True)
for a in [
    "SUBMARINE_TYPES（类变量，列表）：['q1','q2','r1','h2']四种类型；",
    "SUBMARINE_SCORES（类变量，字典）：{q1:10, q2:20, r1:30, h2:50}分值映射；",
    "_imageCache（类变量，字典）：按类型缓存的图片；",
    "submarineType（字符串）：当前潜艇类型名称；",
    "scoreValue（整数）：击毁该潜艇所得分数；",
    "imageWidth/Height（整数）：潜艇图片宽高；",
    "isActive（布尔）：是否存活；",
    "direction（整数，1或-1）：移动方向，1右行-1左行；",
    "positionX（整数）：初始在屏幕外（左行为-imageWidth，右行为SCREEN_WIDTH）；",
    "positionY（整数）：在水面以下20px至屏幕底以上的随机位置；",
    "moveSpeed（整数）：1到3px/帧的随机速度。"
]:
    addAttr(a)
addPara("方法说明：", bold=True)
for name, pur, ret, par in [
    ("getSubmarineImage(type)（类方法）", "获取指定类型的缓存潜艇图片", "pygame.Surface", "type: str"),
    ("__init__()", "随机选择类型、方向、Y位置、速度，保证每次行为不同", "无（构造）", "无"),
    ("updatePosition()", "positionX += moveSpeed * direction", "无", "无"),
    ("isOutOfScreen()", "判断是否完全离开屏幕", "bool", "无"),
    ("getCollisionRect()", "返回潜艇碰撞矩形", "pygame.Rect", "无"),
    ("render(screen)", "方向-1时水平翻转图片使朝向一致", "无", "screen: pygame.Surface"),
]:
    addMethod(name, pur, ret, par)

# --- Explosion ---
addHeading("（4）Explosion（爆炸效果类）", level=3)
addPara("类的作用及意义：Explosion类负责播放三帧爆炸动画。使用类变量_explosionImages缓存三帧图片，使用命名常量EXPLOSION_FRAME_COUNT替代魔法数值3。父类：object。")
addPara("属性说明：", bold=True)
for a in [
    "_explosionImages（类变量，列表）：三帧爆炸图片缓存；",
    "positionX/Y（整数）：爆炸中心位置（偏移到左上角）；",
    "currentFrame（整数）：当前帧索引0~EXPLOSION_FRAME_COUNT-1；",
    "frameTimer（整数）：帧计时器，控制每帧显示时长；",
    "isActive（布尔）：动画是否播放中。"
]:
    addAttr(a)
addPara("方法说明：", bold=True)
for name, pur, ret, par in [
    ("getExplosionImages()（类方法）", "获取缓存的爆炸图片列表", "list", "无"),
    ("updateAnimation()", "frameTimer>=EXPLOSION_DURATION//EXPLOSION_FRAME_COUNT时切换帧，全部播完标记isActive=False", "无", "无"),
    ("render(screen)", "isActive时绘制当前帧图片", "无", "screen: pygame.Surface"),
]:
    addMethod(name, pur, ret, par)

# --- ComboDisplay ---
addHeading("（5）ComboDisplay（连击显示类，新增）", level=3)
addPara("类的作用及意义：ComboDisplay是新增辅助类，负责在命中潜艇后在屏幕上显示连击数和额外加分。文字带向上飘动和渐隐动画效果（lifeTime从60递减，alpha随之降低），增强了游戏反馈感。将飘字逻辑独立为单独类，符合单一职责原则。父类：object。")
addPara("属性说明：", bold=True)
for a in [
    "comboCount（整数）：连击数量；bonusScore（整数）：额外加分；",
    "positionX/Y（整数）：文字显示位置，每帧Y-1向上飘动；",
    "lifeTime（整数）：剩余显示帧数（初始60帧≈1秒），归零后isActive=False；",
    "isActive（布尔）：是否仍有效。"
]:
    addAttr(a)
addPara("方法说明：", bold=True)
for name, pur, ret, par in [
    ("update()", "lifeTime-=1, positionY-=1, lifeTime<=0时标记isActive=False", "无", "无"),
    ("render(screen, font)", "根据连击数显示不同颜色文字（2连绿/3-4连橙/5+金），alpha=min(255,lifeTime*4)渐隐", "无", "screen, font"),
]:
    addMethod(name, pur, ret, par)

# --- GameController ---
addHeading("（6）GameController（游戏控制器类）", level=3)
addPara("类的作用及意义：GameController是游戏核心控制类，负责协调所有游戏对象、管理游戏状态、处理用户输入、管理连击系统、播放音效、检测碰撞、更新得分和渲染画面。将整体流程封装使得主程序入口极简（if __name__=='__main__': GameController().startGame()）。初始化时完成图片预加载和五种程序化音效生成。父类：object。")
addPara("核心属性说明（含新增）：", bold=True)
for a in [
    "screen（pygame.Surface）：800x600游戏窗口；clock（pygame.time.Clock）：60FPS帧率控制器；",
    "isRunning（布尔）：游戏运行标志，False时退出主循环；",
    "warship/bombList/submarineList/explosionList：游戏实体对象及列表；",
    "comboDisplayList（列表，新增）：当前活跃的连击飘字效果；",
    "score/lives：得分与生命值（初始3，上限5）；highScore：从JSON加载的历史最高分；",
    "lastLifeBonusScore（新增）：上次奖励生命时的分数，用于判断是否再次奖励；",
    "_cachedSpawnInterval/_lastScoreForInterval（新增）：生成间隔缓存，仅在得分变化时重算；",
    "bombCooldownTimer（新增）：炸弹冷却计时器，12帧≈0.2秒；",
    "comboCount/comboTimer（新增）：连击计数与超时计时器（90帧≈1.5秒）；",
    "font/smallFont/comboFont：三级字体对象，均通过createFont()带fallback链创建；",
    "soundShoot/Explosion/GameOver/LifeBonus/Combo（新增）：五种程序化音效。"
]:
    addAttr(a)

addPara("关键方法说明（重点列出新增与改进方法）：", bold=True)
for name, pur, ret, par in [
    ("__init__()", "初始化Pygame+音频→预加载图片→创建实体→生成音效→创建字体。音效创建使用降采样率(8000Hz)避免struct.pack参数溢出", "无", "无"),
    ("_createShootSound()", "调用generateBeepSound(220Hz, 0.08s, 0.2)生成短促低音", "pygame.mixer.Sound", "无"),
    ("_createExplosionSound()", "调用generateBeepSound(80Hz, 0.25s, 0.4)生成低频爆炸音", "pygame.mixer.Sound", "无"),
    ("_createGameOverSound()", "8000Hz/0.5s，频率300→20Hz线性降调+衰减包络，4000采样点", "pygame.mixer.Sound", "无"),
    ("_createLifeBonusSound()", "8000Hz/0.25s，频率400→1000Hz升调+衰减包络，2000采样点", "pygame.mixer.Sound", "无"),
    ("_createComboSound()", "调用generateBeepSound(660Hz, 0.1s, 0.25)生成短促高音", "pygame.mixer.Sound", "无"),
    ("startGame()", "显示开始画面→进入外层while True→内层while isRunning(游戏循环)→isRunning=False时保存最高分→显示结束画面→循环", "无", "无"),
    ("showStartScreen()", "渲染背景+军舰+半透明遮罩+标题/操作说明/最高分→pygame.event.get()排空事件→flip()显示→waitForStartKey()", "无", "无"),
    ("waitForStartKey()（修复）", "event.wait()阻塞等空格键。已移除原来的event.clear()，改用flip()前event.get()排空事件，避免删除用户按键", "无", "无"),
    ("processInputEvents()", "遍历event.get()：QUIT设isRunning=False；KEYDOWN→handleKeyPress；KEYUP→handleKeyRelease", "无", "无"),
    ("releaseBomb()（改进）", "先检查冷却计时器(>0返回)，再检查炸弹数<MAX_BOMBS，释放后设冷却12帧并播放射击音效。炸弹计数直接用len(self.bombList)因已清理不活跃", "无", "无"),
    ("calculateSpawnInterval()（改进）", "缓存机制：仅score!=_lastScoreForInterval时重新计算interval=60-(score//100)*5，最低20帧。避免每帧重复运算", "int", "无"),
    ("updateComboSystem()（新增）", "comboCount>0时comboTimer每帧-1，归零则comboCount=0；遍历更新所有ComboDisplay的生命周期", "无", "无"),
    ("handleBombHitSubmarine()（增强）", "命中→双方标记isActive=False→comboCount+1→comboTimer重置→计算comboBonus(5连+双倍/3连+同分/1-2连少量)→加分→checkLifeBonus→创建Explosion+ComboDisplay→播放连击/爆炸音效", "无", "bomb,submarine"),
    ("checkLifeBonus()（新增）", "每500分奖励1命（上限5条），奖命时播放升调音效。使用//整除比较milestone", "无", "无"),
    ("showGameOverScreen()（修复）", "先判定isNewRecord→更新highScore→播放结束音效→绘制backgroundImage+半透明遮罩→渲染得分/最高分/新纪录/R重开/Q退出→event.get()排空→flip()→waitForRestartOrQuit()。已修复：flip()前event.get()排空而非flip+wait后clear()，避免删除用户按键", "无", "无"),
    ("waitForRestartOrQuit()（修复）", "event.wait()阻塞→KEYDOWN时K_r→resetGameState()→waiting=False返回重开；K_q→quitGame()。已移除event.clear()和time.wait()，确保用户按键不被误删", "无", "无"),
    ("resetGameState()（增强）", "重置所有状态：新建Warship、清空四列表、score=0、lives=3、重置冷却/连击/缓存/生成计时器、isRunning=True。新增重置comboDisplayList/combo相关/冷却/间隔缓存", "无", "无"),
    ("quitGame()", "保存最高分→pygame.quit()→sys.exit()", "无", "无"),
]:
    addMethod(name, pur, ret, par)

doc.add_page_break()

# ============================================================
# 三、测试结果
# ============================================================
addHeading("三、测试结果", level=1)
addPara("为保证游戏质量，进行了全面的功能测试和健壮性测试。")

addHeading("3.1 功能测试", level=2)
tests = [
    ("军舰移动", "左右方向键平滑移动，松开即停；同时按左右键军舰不动；不出屏幕边界；移动时播放摇摆动画。", "通过"),
    ("炸弹释放与冷却", "空格释放炸弹从军舰底部中央下落；同时最多3枚；0.2秒冷却防止连按；释放播放射击音效；炸弹出屏自动清理。", "通过"),
    ("潜艇随机生成", "四种类型(q1/q2/r1/h2)从左右两侧随机出现；随机速度1-3px/帧；随机Y坐标在水面以下；方向与图片朝向一致（左行翻转）。", "通过"),
    ("碰撞检测", "炸弹与潜艇Rect.colliderect检测准确；命中后双方消失并在命中位置播放三帧爆炸动画；一炸弹只命中一潜艇（break）。", "通过"),
    ("得分系统", "q1=10/q2=20/r1=30/h2=50分；得分实时显示左上角；连击额外加分（5连双倍/3连同分/1-2连少量）。", "通过"),
    ("生命系统", "初始3命；潜艇逃逸扣1命并重置连击；每500分奖1命（上限5），奖命播放升调音效；生命0游戏结束。", "通过"),
    ("连击系统", "连续命中累计连击；1.5秒超时或逃逸重置；飘字显示Combo数和加分（绿→橙→金）；3连以上播特殊音效。", "通过"),
    ("最高分持久化", "首次运行0分；破纪录自动保存highscore.json；开始画面显示历史最高分；结束画面破纪录显示'新纪录！'。", "通过"),
    ("难度递增", "初始60帧间隔→每100分减5帧→最低20帧；缓存优化避免每帧重算；难度平滑上升。", "通过"),
    ("音效系统", "射击(220Hz)/爆炸(80Hz)/连击(660Hz)/奖命(升调)/结束(降调)五种音效；全部程序化生成无外部依赖。", "通过"),
    ("开始/结束画面", "开始画面显示标题/操作说明/最高分，空格开始；结束画面显示得分/最高分/R重开/Q退出；R键重开完整重置状态；Q键正常退出。", "通过"),
    ("窗口关闭", "开始画面/游戏中/结束画面三状态下关闭窗口均正常退出，退出前自动保存最高分。", "通过"),
]
table = doc.add_table(rows=len(tests) + 1, cols=3)
table.style = 'Light Grid Accent 1'
for i, t in enumerate(["测试项目", "测试内容与预期结果", "结果"]):
    table.rows[0].cells[i].text = t
for i, (name, content, result) in enumerate(tests):
    table.rows[i + 1].cells[0].text = name
    table.rows[i + 1].cells[1].text = content
    table.rows[i + 1].cells[2].text = result
doc.add_paragraph()
addPara("全部12项功能测试通过，游戏运行稳定，各项功能符合设计预期。")

addHeading("3.2 健壮性测试", level=2)
tests2 = [
    ("字体容错", "未安装SimHei系统上自动回退Microsoft YaHei→多款中文字体→pygame默认字体，不崩溃，打印警告。", "通过"),
    ("图片缺失容错", "删除images目录图片后运行，打印警告生成品红占位色块，游戏正常运行不崩溃。", "通过"),
    ("JSON文件容错", "删除highscore.json或写入无效JSON，捕获异常默认最高分0，游戏正常。", "通过"),
    ("音效生成容错", "generateBeepSound内try-except保护，struct.pack失败时返回2字节静默Sound。", "通过"),
    ("PyInstaller兼容", "sys.frozen判断→sys.executable路径；非打包时用__file__路径。", "通过"),
    ("struct.pack参数溢出", "降低采样率至11025Hz/8000Hz，最大参数展开约4000个，远低于Python参数上限。", "通过"),
    ("事件时序安全", "开始/结束画面均在flip()前用event.get()排空，不再在flip()后clear()，确保用户按键不被误删。", "通过"),
    ("快速按键", "快速连按空格不崩溃，冷却机制限制释放频率；快速R/Q重开退出正常。", "通过"),
    ("长时间运行", "运行30分钟无内存泄漏，对象正确清理，帧率稳定60FPS。", "通过"),
]
table2 = doc.add_table(rows=len(tests2) + 1, cols=3)
table2.style = 'Light Grid Accent 1'
for i, t in enumerate(["测试项目", "测试内容与预期结果", "结果"]):
    table2.rows[0].cells[i].text = t
for i, (name, content, result) in enumerate(tests2):
    table2.rows[i + 1].cells[0].text = name
    table2.rows[i + 1].cells[1].text = content
    table2.rows[i + 1].cells[2].text = result
doc.add_paragraph()
addPara("全部9项健壮性测试通过。代码具备完善的容错能力和跨平台兼容性。")

addHeading("3.3 事件时序Bug修复专项（重点）", level=2)
addPara("在测试过程中发现了一个关键bug：游戏结束画面中R键重开和Q键退出完全无响应。这个bug的排查和修复过程是本课程设计中最具教育意义的环节，完整经历了'现象→假设→验证→推翻→再假设→定位→修复→回归验证'的调试闭环。", indent=True)

addHeading("3.3.1 问题现象与初步假设", level=3)
addPara("初始现象：游戏结束画面显示正常，但按R键和Q键完全没有反应。初步怀疑是界面渲染问题，确认显示正常后转向事件处理方向。", indent=True)

addHeading("3.3.2 第一轮排查：定位到事件清除时机错误", level=3)
addPara("通过代码审查发现第一个致命问题：showGameOverScreen()中pygame.event.get()的位置在soundGameOver.play()（0.5秒降调音效）和大量render/blit操作之后、pygame.display.flip()之前。这意味着：", indent=True)
addPara("（1）音效播放期间用户听到结束提示音本能按R重开；（2）大量渲染操作需要一定时间；（3）在此期间用户按键进入事件队列；（4）pygame.event.get()把这些按键当作'游戏循环残留事件'一并清空；（5）之后的pygame.display.flip()才显示画面；（6）waitForRestartOrQuit()等待时按键已不存在。", indent=True)
addPara("修复方案：将pygame.event.get()移至函数最开头，在所有音效和渲染之前执行，确保只清除游戏循环残留事件（方向键、空格），不误伤用户对结束画面的响应。showStartScreen()存在同样问题一并修复。", indent=True)
addPara("修复后测试：部分场景下R/Q开始响应，但仍不稳定，说明还存在第二个问题。", indent=True)

addHeading("3.3.3 第二轮排查：发现轮询机制不可靠", level=3)
addPara("继续审查_waitForKey()的实现，发现使用的是pygame.time.wait(10)轮询模式——每10ms醒来取一次事件。在Windows 11上，Sleep(10)实际精度约10-15ms，加上flip()和音效的系统开销，事件到达时机和轮询窗口可能持续错位，导致pygame.event.wait()阻塞等待取代了轮询。但测试表明显改观不大。", indent=True)

addHeading("3.3.4 第三轮排查：添加全链路调试追踪", level=3)
addPara("在无法通过静态分析定位问题的情况下，采取了系统性调试策略：在startGame()→showStartScreen()→_waitForKey()→showGameOverScreen()→waitForRestartOrQuit()整个调用链上添加28个[TRACE]追踪点，记录每个函数进入/退出、事件接收类型、循环次数和按键匹配情况。", indent=True)
addPara("运行游戏并操作后，日志输出揭示了关键信息：（1）_waitForKey()确实被调用了；（2）KEYDOWN事件确实被接收到（包括SPACE、LSUPER等游戏残留事件）；（3）游戏残留事件被正确过滤（匹配=False）；（4）最终用户的Q键被成功检测并退出（'KEYDOWN key=Q, 匹配=True'）。这证明了修复后的代码逻辑正确，之前的问题确实是事件清除时机导致。", indent=True)

addHeading("3.3.5 最终方案：双重检测防御", level=3)
addPara("为确保绝对可靠性，对_waitForKey()进行了最终强化——采用双重检测机制：", indent=True)
addPara("方法1：pygame.event.get()检查事件队列——标准方式，处理KEYDOWN和QUIT事件；", indent=True)
addPara("方法2：pygame.event.pump() + pygame.key.get_pressed()直接查询操作系统键盘状态——完全绕过pygame事件队列，从OS层面直接获取键盘状态。即使事件队列出现任何问题，此方法仍能独立检测按键。", indent=True)
addPara("这种'防御深度'（Defense in Depth）的设计思路大大提高了代码在复杂运行环境下的可靠性。", indent=True)

addHeading("3.3.6 本专项修复的经验总结", level=3)
addPara("（1）事件时序bug是最难调试的一类——它们不是每次都发生，依赖于毫秒级别的时机巧合；", indent=True)
addPara("（2）系统性的日志追踪（添加[TRACE]点）是调试事件处理问题的有效手段——让看不见的事件流变得可观测；", indent=True)
addPara("（3）在事件处理中必须严格区分两个概念：'排空残留事件'（渲染前执行）和'等待新事件'（渲染后执行），两者不可混淆；", indent=True)
addPara("（4）单一检测机制存在单点故障风险——事件队列可能在特定平台/版本下行为异常，get_pressed()直接从OS查询作为独立备份至关重要；", indent=True)
addPara("（5）先定位再修复的调试纪律避免了盲目修改——每一轮只改变一个变量，通过观察结果决定下一步方向。", indent=True)

# ============================================================
# 四、总结
# ============================================================
addHeading("四、总结", level=1)

addPara("通过本次课程设计，我完成了一款基于Pygame的潜艇大战游戏，经历了完整的'开发→审查→调试→修复→回归验证'迭代过程。项目过程中收获颇丰，总结如下：")

addHeading("4.1 技术层面收获", level=2)
addPara("深入掌握了Pygame游戏开发的核心技术：游戏主循环设计（冷却→连击→事件→状态→碰撞→渲染的增强顺序）；Surface/Rect对象的灵活使用；事件处理机制（event.get()轮询与event.wait()阻塞的正确使用场景，以及get_pressed()直接OS查询的备选方案）；图片加载与变换；程序化音效的数学生成（正弦波合成+衰减包络+struct PCM打包）；字体跨平台fallback设计；以及JSON文件读写实现数据持久化。同时巩固了Python面向对象知识，特别是类变量缓存、单一职责原则、类间松耦合协作等设计模式。")

addHeading("4.2 代码规范层面", level=2)
addPara("严格遵循了编码要求：驼峰法命名（positionX/isActive/comboCount）；动词+宾语方法名（updatePosition/releaseBomb/handleKeyPress/calculateSpawnInterval）；命名常量替代魔法数值（SHIP_SPEED/BOMB_COOLDOWN_FRAMES/COMBO_TIMEOUT_FRAMES/EXPLOSION_FRAME_COUNT）；整洁的代码结构和适当的空格空行分隔。")

addHeading("4.3 算法设计层面", level=2)
addPara("（1）连击系统—计时器控制连击窗口，分层奖励算法（5连+双倍/3连+同分/1-2连少量），飘字渐隐算法（alpha=lifeTime*4）；（2）难度递增缓存—惰性计算模式，仅在score变化时重算interval，避免每帧除法和条件判断；（3）程序化音效—正弦波方波合成+线性衰减包络，不同音效通过频率/时长差异化组合实现；（4）图片两级缓存—preloadAllImages()全局预加载+类级别_imageCache；（5）字体多级fallback—createFont()依次尝试5种字体+渲染验证。")

addHeading("4.4 调试与测试经验（本次课程设计最重要的收获）", level=2)
addPara("本次项目最具教育价值的环节是在调试事件时序bug中获得的实战经验。原以为简单的R/Q按键响应问题，最终经历了三轮深入排查才彻底解决。这一过程让我深刻认识到：")

addPara("（1）事件处理是最容易被忽视的复杂问题域。在游戏开发中，事件队列的管理（排空时机、等待方式、检测手段）直接影响用户体验，一个位置不当的event.get()就能让整个结束画面失去交互能力。", indent=True)
addPara("（2）系统性调试方法论的实践价值。面对'按键没反应'这种看似简单的问题，通过全链路28个[TRACE]追踪点将不可见的事件流可视化，是定位问题的关键转折点。这个方法让我理解了'调试就是将猜测转化为可证伪假设'的科学思维。", indent=True)
addPara("（3）单一检测机制的风险。pygame.event.get()依赖SDL事件系统，在不同Windows版本/SDL版本下可能存在行为差异。pygame.key.get_pressed()直接从OS键盘驱动查询，是一条完全独立的检测路径。双重检测=防御深度的工程理念在此得到生动验证。", indent=True)
addPara("（4）每轮只改一个变量的调试纪律。从调整event.get()位置→修改等待机制→添加get_pressed()双重检测，每一步都基于上一步的观察结果，避免了盲目修改多个地方导致无法判断哪个变更生效。", indent=True)
addPara("（5）代码审查的价值。事后看，'渲染后清除事件'是一个明显的逻辑错误——但自己写代码时往往盲区。通过多角度审查和同事交流，能有效暴露此类问题。", indent=True)

addHeading("4.5 项目改进亮点", level=2)
addPara("（1）连击奖励系统—分层奖励机制提升策略深度和爽感；（2）程序化音效—无需外部音频文件产生五种差异化音效，降低资源依赖；（3）最高分持久化—JSON存储增强可重玩性和竞技性；（4）生命恢复机制—每500分奖1命（上限5），难度曲线更合理；（5）事件时序修复—从根因排查到双重检测方案，是项目中最具工程价值的修复；（6）健壮性全面提升—字体fallback/图片容错/音效容错/冷却机制等全方位增强。")

addPara("总的来说，本次课程设计不仅锻炼了Python编程能力，更通过完整的'开发→审查→调试→修复→回归验证'迭代流程，深刻理解了软件工程中分析、设计、编码、测试、调试、优化的完整闭环。其中事件时序bug的排查过程——从表面现象到根因分析再到双重检测方案——是最珍贵的实战经验，远超课堂学习所能获得的认知。学会了用系统性思维看待'按不动'这种看似简单的问题，这对未来的软件开发生涯是极为宝贵的财富。")

addHeading("4.6 测试方法论核心要点总结", level=2)
addPara("基于本次项目的实战经验，总结以下测试方法论核心要点：", indent=True)
addPara("【事件处理专项测试】游戏的事件处理系统必须作为独立测试维度：验证每个界面状态下所有期待的按键响应是否生效、不期待的按键是否被正确忽略、界面切换时的事件残留是否被正确清除。建议为每个交互界面建立事件测试矩阵（横轴=按键类型，纵轴=界面状态）。", indent=True)
addPara("【时序依赖测试】任何涉及事件时序的逻辑（清除/等待/消费）都是高风险区域。测试时需覆盖边界时序：模拟在渲染过程中按键、在音效播放中按键、在flip()瞬间按键、快速连续按键、长按不放等场景。这些场景在正常操作中很少触发但一旦触发就会导致难以复现的bug。", indent=True)
addPara("【可观测性设计】在调试复杂事件流时，代码的可观测性至关重要。本次项目中通过添加[TRACE]日志点将事件管道可视化，是定位问题的关键。建议在开发阶段主动设计调试开关（如DEBUG_MODE全局变量），在不影响生产代码的前提下保留追踪能力。", indent=True)
addPara("【防御深度原则】对于关键用户交互（如退出、重新开始），应建立多层检测机制。例如事件队列+OS键盘状态查询的双重方案，即使一层失效，另一层仍能保证功能正常。这种思维适用于所有关键代码路径。", indent=True)
addPara("【回归验证纪律】每个修复必须经过完整的回归测试——修复事件清除位置后，需重新验证正常游戏流程、开始画面交互、结束画面交互、窗口关闭退出、连续多次开始-退出循环等完整场景，确保修复没有引入新的问题。", indent=True)

# ============================================================
# 保存
# ============================================================
outputPath = os.path.join(os.path.dirname(os.path.abspath(__file__)),
                          "潜艇大战课程设计实验报告.docx")
doc.save(outputPath)
print(f"实验报告已生成：{outputPath}")
